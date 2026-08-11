import io
import json
import os
import re
import sqlite3
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from flask import Flask, jsonify, request
from flask_cors import CORS
from PyPDF2 import PdfReader
from docx import Document

try:
    from google import genai
    from google.genai import types
except Exception:
    genai = None
    types = None

try:
    import firebase_admin
    from firebase_admin import credentials, firestore
except Exception:
    firebase_admin = None
    credentials = None
    firestore = None

app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 8 * 1024 * 1024

STOPWORDS = {"the","and","for","with","that","this","from","your","you","are","was","were","have","has","had","will","our","their","they","them","into","about","using","use","user","role","job","work","team","years","year","required","preferred","skills","skill","experience","responsibilities","responsibility","candidate","looking","including","such","more","than","also","can","may","must","should","who","what","where","when","how","all","any","but","not","his","her","its","a","an","of","to","in","on","at","by","as","is","be","or","it","we","i","me","my"}


def _db_path():
    base = Path("/tmp") if os.environ.get("VERCEL") else Path(__file__).resolve().parent.parent
    return base / "resume_analyzer.db"


def init_db():
    try:
        con = sqlite3.connect(_db_path())
        con.execute("CREATE TABLE IF NOT EXISTS analyses (id INTEGER PRIMARY KEY AUTOINCREMENT, created_at TEXT, filename TEXT, score INTEGER, target_role TEXT, result_json TEXT)")
        con.commit()
        con.close()
    except Exception:
        pass


def save_sqlite(filename, score, role, result):
    try:
        con = sqlite3.connect(_db_path())
        con.execute("INSERT INTO analyses(created_at,filename,score,target_role,result_json) VALUES(?,?,?,?,?)", (datetime.now(timezone.utc).isoformat(), filename, score, role, json.dumps(result)))
        con.commit()
        con.close()
    except Exception:
        pass


def save_firestore(filename, score, role, result):
    if not all(os.getenv(k) for k in ["FIREBASE_PROJECT_ID", "FIREBASE_CLIENT_EMAIL", "FIREBASE_PRIVATE_KEY"]):
        return False
    if firebase_admin is None:
        return False
    try:
        if not firebase_admin._apps:
            private_key = os.environ["FIREBASE_PRIVATE_KEY"].replace("\\n", "\n")
            cred = credentials.Certificate({"type": "service_account", "project_id": os.environ["FIREBASE_PROJECT_ID"], "private_key": private_key, "client_email": os.environ["FIREBASE_CLIENT_EMAIL"], "token_uri": "https://oauth2.googleapis.com/token"})
            firebase_admin.initialize_app(cred)
        firestore.client().collection("resume_analyses").add({"createdAt": datetime.now(timezone.utc), "filename": filename, "score": score, "targetRole": role, "result": result})
        return True
    except Exception:
        return False


def extract_pdf(data):
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def extract_docx(data):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(data)
        path = f.name
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def extract_text(file_storage):
    name = (file_storage.filename or "").lower()
    data = file_storage.read()
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".docx"):
        return extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Supported formats are PDF, DOCX, and TXT.")


def words(text):
    return re.findall(r"[a-zA-Z][a-zA-Z0-9+#./-]{1,}", text.lower())


def normalize_terms(text):
    return {w.strip(".,:;()[]{}") for w in words(text) if w not in STOPWORDS and len(w) > 2}


def keyword_overlap(resume, jd):
    r = normalize_terms(resume)
    j = normalize_terms(jd)
    if not j:
        return 0, []
    matched = sorted(r & j)
    return round(100 * len(matched) / max(1, len(j))), matched


def deterministic_score(resume, jd, role):
    low = resume.lower()
    sections = {
        "Contact & Identity": bool(re.search(r"(?:@|phone|linkedin|github|mobile|contact)", low)),
        "Education": any(k in low for k in ["education", "b.sc", "b.e", "b.tech", "bachelor", "degree", "university", "college"]),
        "Projects": any(k in low for k in ["project", "developed", "built", "implemented"]),
        "Skills": any(k in low for k in ["technical skills", "skills", "python", "javascript", "sql"]),
        "Testing": any(k in low for k in ["test", "testing", "qa", "debug", "postman"]),
        "Action-oriented content": sum(1 for k in ["built", "developed", "implemented", "designed", "integrated", "tested", "optimized", "automated"] if k in low) >= 3,
        "Quantified evidence": bool(re.search(r"\b\d+(?:\.\d+)?\s*(?:%|ms|x|users|projects|years|months|apis|tests|gestures|features)\b", low)),
        "ATS readability": not any(x in low for x in ["references available", "date of birth", "marital status"]),
    }
    completeness = round(100 * sum(sections.values()) / len(sections))
    if jd.strip():
        kw, matched = keyword_overlap(resume, jd)
    else:
        tech_hits = sum(1 for k in ["python", "flask", "javascript", "sql", "rest", "api", "git", "react", "next.js", "firebase", "testing", "postman", "html", "css"] if k in low)
        kw, matched = min(100, round(tech_hits / 13 * 100)), []
    score = round(0.45 * completeness + 0.35 * kw + 0.20 * (100 if sections["Action-oriented content"] else 50))
    score = max(0, min(100, score))
    greens, reds, recs = [], [], []
    if sections["Contact & Identity"]:
        greens.append("Contact details and professional profile links are present.")
    else:
        reds.append("Contact or professional profile information appears incomplete.")
    if sections["Projects"]:
        greens.append("Projects use action-oriented engineering language.")
    else:
        reds.append("Project evidence is too limited or difficult to identify.")
    if sections["Skills"]:
        greens.append("Technical skills are explicitly grouped and searchable by ATS systems.")
    else:
        reds.append("Technical skills are not clearly separated into a skills section.")
    if sections["Testing"]:
        greens.append("Testing and debugging evidence is visible, which strengthens developer and QA positioning.")
    else:
        recs.append("Add concrete testing or debugging evidence if you performed it.")
    if not sections["Quantified evidence"]:
        reds.append("The resume has limited measurable outcomes or metrics.")
        recs.append("Add truthful metrics such as supported features, test coverage, latency, accuracy, or project scale where available.")
    if jd.strip() and kw < 60:
        reds.append("Job-description keyword alignment is currently below a strong ATS match level.")
        recs.append("Mirror relevant terminology from the job description only when it accurately reflects your skills.")
    if jd.strip():
        recs.append("Prioritize matched job-specific skills in the summary, skills section, and most relevant project bullets.")
    return score, greens[:8], reds[:8], recs[:10], matched[:60], sections


SCHEMA = {"type": "object", "properties": {"summary": {"type": "string"}, "green_flags": {"type": "array", "items": {"type": "string"}}, "red_flags": {"type": "array", "items": {"type": "string"}}, "recommendations": {"type": "array", "items": {"type": "string"}}, "strengths": {"type": "array", "items": {"type": "string"}}, "risks": {"type": "array", "items": {"type": "string"}}, "ats_keywords": {"type": "array", "items": {"type": "string"}}, "role_fit": {"type": "string"}, "score_adjustment": {"type": "integer"}}, "required": ["summary", "green_flags", "red_flags", "recommendations", "strengths", "risks", "ats_keywords", "role_fit", "score_adjustment"]}


def gemini_review(resume, jd, role):
    key = os.getenv("GEMINI_API_KEY")
    if not key or genai is None or types is None:
        return None
    client = genai.Client(api_key=key)
    prompt = f"""You are a senior technical recruiter and ATS specialist. Analyze this resume for an entry-level candidate. Be evidence-based: never invent experience, certifications, metrics, technologies, or achievements. Separate objective red flags from optional improvements. Assess fit for: {role or 'Software Developer / QA / IT'}. Job description (may be empty):\n{jd[:12000]}\n\nRESUME:\n{resume[:24000]}\n\nReturn structured JSON only. Score adjustment must be an integer from -8 to +8 and should only adjust the deterministic score when the text supports it. Focus on ATS keyword alignment, clarity, credibility, technical specificity, impact, project quality, testing evidence, and missing essentials."""
    try:
        resp = client.models.generate_content(model=os.getenv("GEMINI_MODEL", "gemini-2.0-flash"), contents=prompt, config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=SCHEMA, temperature=0.2, max_output_tokens=1800))
        return json.loads(resp.text)
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/health")
@app.get("/health")
def health():
    return jsonify({"status": "ok", "gemini_configured": bool(os.getenv("GEMINI_API_KEY")), "firebase_configured": all(os.getenv(k) for k in ["FIREBASE_PROJECT_ID", "FIREBASE_CLIENT_EMAIL", "FIREBASE_PRIVATE_KEY"])})


@app.post("/api/analyze")
@app.post("/analyze")
def analyze():
    try:
        if "resume" not in request.files:
            return jsonify({"error": "Upload a resume file first."}), 400
        file = request.files["resume"]
        if not file.filename:
            return jsonify({"error": "The uploaded file has no filename."}), 400
        text = extract_text(file)
        if len(text.strip()) < 120:
            return jsonify({"error": "Not enough readable text was extracted. Try a text-based PDF or DOCX."}), 422
        jd = request.form.get("job_description", "").strip()
        role = request.form.get("target_role", "Software Developer / QA / IT").strip()
        score, greens, reds, recs, matched, sections = deterministic_score(text, jd, role)
        ai = gemini_review(text, jd, role)
        ai_error = None
        if ai and "error" not in ai:
            score = max(0, min(100, score + int(ai.get("score_adjustment", 0))))
            greens = list(dict.fromkeys(greens + ai.get("green_flags", [])))[:10]
            reds = list(dict.fromkeys(reds + ai.get("red_flags", [])))[:10]
            recs = list(dict.fromkeys(recs + ai.get("recommendations", [])))[:12]
            matched = list(dict.fromkeys(matched + ai.get("ats_keywords", [])))[:60]
        elif ai and "error" in ai:
            ai_error = "Gemini enhancement was unavailable for this run; the explainable deterministic ATS analysis was still completed."
        result = {"score": score, "score_label": "Excellent" if score >= 85 else "Strong" if score >= 75 else "Needs Improvement" if score >= 60 else "Weak", "summary": (ai or {}).get("summary") if ai and "error" not in ai else f"The resume received a {score}/100 baseline score based on structure, technical evidence, project content, testing evidence, and job-description alignment.", "green_flags": greens, "red_flags": reds, "recommendations": recs, "strengths": (ai or {}).get("strengths", []) if ai and "error" not in ai else [], "risks": (ai or {}).get("risks", []) if ai and "error" not in ai else [], "matched_keywords": matched, "role_fit": (ai or {}).get("role_fit", "Baseline role-fit analysis completed.") if ai and "error" not in ai else "Baseline role-fit analysis completed.", "sections": sections, "filename": file.filename, "target_role": role, "job_description_provided": bool(jd), "ai_enhanced": bool(ai and "error" not in ai), "ai_notice": ai_error, "word_count": len(words(text))}
        save_sqlite(file.filename, score, role, result)
        save_firestore(file.filename, score, role, result)
        return jsonify(result)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": "Analysis failed. Please try again with a text-based PDF or DOCX.", "detail": str(e)}), 500


init_db()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "5000")), debug=True)
